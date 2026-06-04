"""
Construye TFM_Entrega2.docx con formato UNIR correcto.

Estrategia:
- Abre v1.1 como plantilla (preserva todos los estilos, márgenes, encabezados y pies).
- Borra el contenido del cuerpo conservando los section properties.
- Parsea cada capítulo .tex e inserta párrafos con los estilos correctos.
- Incrusta las figuras PNG centradas con caption encima.
- Reconstruye las tablas como tablas Word con la cabecera en negrita.
"""

from __future__ import annotations

import copy
import os
import re
import sys
from pathlib import Path

ROOT = Path(__file__).parent
MEMORIA = ROOT / "Memoria"
TEMPLATE = MEMORIA / "Versiones y Entregas" / "v1.1_Entrega01_CamiloRico_JuanBlanco_JoseMachado-Corregida.docx"
OUTPUT = MEMORIA / "TFM_Entrega2.docx"

import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color UNIR ───────────────────────────────────────────────────────────────
UNIR_BLUE = RGBColor(0x00, 0x98, 0xCD)

# ── Mapa de citas (key → (autores_texto, autores_paren, año)) ─────────────────
CITAS = {
    "Armstrong2001":   ("Armstrong", "Armstrong", "2001"),
    "Bergmeir2012":    ("Bergmeir y Hyndman", "Bergmeir & Hyndman", "2012"),
    "Hyndman2021":     ("Hyndman y Athanasopoulos", "Hyndman & Athanasopoulos", "2021"),
    "Januschowski2020":("Januschowski et al.", "Januschowski et al.", "2020"),
    "LimZohren2021":   ("Lim y Zohren", "Lim & Zohren", "2021"),
    "LimTFT2021":      ("Lim et al.", "Lim et al.", "2021"),
    "Makridakis2022":  ("Makridakis et al.", "Makridakis et al.", "2022"),
    "Makridakis2025M6":("Makridakis et al.", "Makridakis et al.", "2025"),
    "Makridakis2020M4":("Makridakis et al.", "Makridakis et al.", "2020"),
    "Petropoulos2022": ("Petropoulos et al.", "Petropoulos et al.", "2022"),
    "Oreshkin2020":    ("Oreshkin et al.", "Oreshkin et al.", "2020"),
    "Ansari2024":      ("Ansari et al.", "Ansari et al.", "2024"),
    "Rahimikia2025":   ("Rahimikia et al.", "Rahimikia et al.", "2025"),
    "Zhu2025FinCast":  ("Zhu et al.", "Zhu et al.", "2025"),
    "Cerqueira2020":   ("Cerqueira et al.", "Cerqueira et al.", "2020"),
    "Cerqueira2022":   ("Cerqueira et al.", "Cerqueira et al.", "2022"),
    "Fildes2022":      ("Fildes et al.", "Fildes et al.", "2022"),
    "Kolassa2022":     ("Kolassa", "Kolassa", "2022"),
    "Zhou2021Informer":("Zhou et al.", "Zhou et al.", "2021"),
}

# ── Referencias APA completas para la sección de bibliografía ─────────────────
REFERENCIAS_APA = [
    ("Armstrong2001",
     "Armstrong, J. S. (2001). Principles of Forecasting: A Handbook for Researchers "
     "and Practitioners. Kluwer Academic Publishers. https://doi.org/10.1007/978-0-306-47630-3"),
    ("Bergmeir2012",
     "Bergmeir, C., & Hyndman, R. J. (2012). On the use of cross-validation for time "
     "series predictor evaluation. Information Sciences, 191, 192–203. "
     "https://doi.org/10.1016/j.ins.2011.12.028"),
    ("Cerqueira2020",
     "Cerqueira, V., Torgo, L., & Mozetič, I. (2020). Evaluating time series forecasting "
     "models: An empirical study on performance estimation methods. Machine Learning, 109, "
     "1997–2028. https://doi.org/10.1007/s10994-020-05910-7"),
    ("Cerqueira2022",
     "Cerqueira, V., Torgo, L., & Soares, C. (2022). A case study comparing machine "
     "learning with statistical methods for time series forecasting: size matters. "
     "Journal of Intelligent Information Systems, 59(2), 415–433. "
     "https://doi.org/10.1007/s10844-022-00713-9"),
    ("Fildes2022",
     "Fildes, R., Ma, S., & Kolassa, S. (2022). Retail forecasting: Research and practice. "
     "International Journal of Forecasting, 38(4), 1283–1318. "
     "https://doi.org/10.1016/j.ijforecast.2019.06.004"),
    ("Hyndman2021",
     "Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and Practice "
     "(3.ª ed.). OTexts. https://otexts.com/fpp3/"),
    ("Januschowski2020",
     "Januschowski, T., Wang, Y., Torkkola, K., Erkkilä, T., Hasson, H., & Gasthaus, J. "
     "(2020). Forecasting with trees. International Journal of Forecasting, 36(1), 167–173. "
     "https://doi.org/10.1016/j.ijforecast.2019.10.004"),
    ("Kolassa2022",
     "Kolassa, S. (2022). Why the «Best» Point Forecast Depends on the Error or Accuracy "
     "Measure. En M. Gilliland, L. Tashman y U. Sglavo (Eds.), Business Forecasting: "
     "Practical Problems and Solutions (pp. 113–122). Wiley."),
    ("LimZohren2021",
     "Lim, B., & Zohren, S. (2021). Time-series forecasting with deep learning: a survey. "
     "Philosophical Transactions of the Royal Society A, 379(2194), 20200209. "
     "https://doi.org/10.1098/rsta.2020.0209"),
    ("LimTFT2021",
     "Lim, B., Arık, S. Ö., Loeff, N., & Pfister, T. (2021). Temporal Fusion Transformers "
     "for interpretable multi-horizon time series forecasting. International Journal of "
     "Forecasting, 37(4), 1748–1764. https://doi.org/10.1016/j.ijforecast.2021.03.012"),
    ("Makridakis2020M4",
     "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2020). The M4 Competition: "
     "100,000 time series and 61 forecasting methods. International Journal of Forecasting, "
     "36(1), 54–74. https://doi.org/10.1016/j.ijforecast.2019.04.014"),
    ("Makridakis2022",
     "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022). M5 accuracy competition: "
     "Results, findings, and conclusions. International Journal of Forecasting, 38(4), "
     "1346–1364. https://doi.org/10.1016/j.ijforecast.2021.11.013"),
    ("Makridakis2025M6",
     "Makridakis, S., Spiliotis, E., Hollyman, R., Petropoulos, F., Swanson, N., & Gaba, A. "
     "(2025). The M6 forecasting competition: Bridging the gap between forecasting and "
     "investment decisions. International Journal of Forecasting, 41(4), 1315–1354. "
     "https://doi.org/10.1016/j.ijforecast.2024.11.002"),
    ("Ansari2024",
     "Ansari, A. F., Stella, L., Turkmen, C., Zhang, X., Mercado, P., Shen, H., "
     "Shchur, O., Rangapuram, S. S., Pineda Arango, S., Kapoor, S., Zschiegner, J., "
     "Maddix, D. C., Mahoney, M. W., Januschowski, T., de Jong, V., "
     "Bohlke-Schneider, M., & Wang, Y. (2024). Chronos: Learning the Language of "
     "Time Series. arXiv:2403.07815. https://arxiv.org/abs/2403.07815"),
    ("Oreshkin2020",
     "Oreshkin, B. N., Carpov, D., Chapados, N., & Bengio, Y. (2020). N-BEATS: Neural "
     "basis expansion analysis for interpretable time series forecasting. En Proceedings "
     "of the 8th International Conference on Learning Representations (ICLR 2020). "
     "https://openreview.net/forum?id=r1ecqn4YwB"),
    ("Petropoulos2022",
     "Petropoulos, F., Apiletti, D., Assimakopoulos, V., Babai, M. Z., Barrow, D. K., "
     "Ben Taieb, S., Bergmeir, C., Bessa, R. J., Bijak, J., Boylan, J. E., Browell, J., "
     "y otros. (2022). Forecasting: theory and practice. International Journal of "
     "Forecasting, 38(3), 705–871. https://doi.org/10.1016/j.ijforecast.2021.11.001"),
    ("Rahimikia2025",
     "Rahimikia, E., Ni, H., & Wang, W. (2025). Re(Visiting) Time Series Foundation "
     "Models in Finance. arXiv:2511.18578. https://arxiv.org/abs/2511.18578"),
    ("Zhu2025FinCast",
     "Zhu, Z., Chen, H., Qu, Q., & Chung, V. (2025). FinCast: A Foundation Model for "
     "Financial Time-Series Forecasting. En Proceedings of the 34th ACM International "
     "Conference on Information and Knowledge Management (CIKM 2025). "
     "https://doi.org/10.1145/3746252.3761261"),
    ("Zhou2021Informer",
     "Zhou, H., Zhang, S., Peng, J., Zhang, S., Li, J., Xiong, H., & Zhang, W. (2021). "
     "Informer: Beyond Efficient Transformer for Long Sequence Time-Series Forecasting. "
     "Proceedings of the AAAI Conference on Artificial Intelligence, 35(12), 11106–11115. "
     "https://doi.org/10.1609/aaai.v35i12.17325"),
]


# ═══════════════════════════════════════════════════════════════════════════════
# LATEX PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def fmt_parencite(keys: str) -> str:
    """Convierte claves de cita a formato parentético APA en español."""
    parts = []
    for k in keys.split(","):
        k = k.strip()
        if k in CITAS:
            a, _, y = CITAS[k]
            parts.append(f"{a}, {y}")
    return "(" + "; ".join(parts) + ")" if parts else f"({keys})"


def fmt_textcite(keys: str) -> str:
    """Convierte claves de cita a formato narrativo APA en español."""
    parts = []
    for k in keys.split(","):
        k = k.strip()
        if k in CITAS:
            a, _, y = CITAS[k]
            parts.append(f"{a} ({y})")
    return "; ".join(parts) if parts else keys


def clean_inline(text: str) -> str:
    """Aplica limpieza de comandos LaTeX inline al texto plano.

    Orden critico:
    1. Citas y referencias estructurales.
    2. Caracteres especiales LaTeX (incluido \\% escapeado).
    3. Comandos layout SIN argumento ANTES de los de formato, usando
       negative lookahead (?![a-zA-Z]) para no fusionarse con texto pegado.
    4. Formato de texto (textbf, textit...) en multiples pasadas para
       manejar anidamiento.
    5. Limpieza final de residuos.
    """
    # 1. Citas
    text = re.sub(r"\\parencite\{([^}]+)\}", lambda m: fmt_parencite(m.group(1)), text)
    text = re.sub(r"\\textcite\{([^}]+)\}", lambda m: fmt_textcite(m.group(1)), text)
    text = re.sub(r"\\cite\{([^}]+)\}", lambda m: fmt_parencite(m.group(1)), text)

    # 2. Referencias estructurales
    text = re.sub(r"\\(?:ref|label|pageref)\{[^}]+\}", "", text)
    text = re.sub(r"\\url\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\href\{[^}]+\}\{([^}]+)\}", r"\1", text)

    # 3. Caracteres especiales LaTeX
    text = text.replace(r"\&", "&")
    text = text.replace(r"\%", "%")    # debe ser ANTES de cualquier regex con %
    text = text.replace(r"\$", "$")
    text = text.replace(r"\_", "_")
    text = text.replace(r"\#", "#")
    text = text.replace(r"\,", " ")
    text = text.replace("~", " ")
    text = text.replace("---", "—")
    text = text.replace("--", "–")
    text = text.replace("``", "“").replace("''", "”")
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    text = text.replace(r"\\", " ")    # LaTeX newline -> space

    # 4. Matematicas simples
    text = re.sub(r"\$\\pm\$", "±", text)
    text = re.sub(r"\$\\geq\$", "≥", text)
    text = re.sub(r"\$\\leq\$", "≤", text)
    text = re.sub(r"\$\\Delta\$", "Δ", text)
    text = re.sub(r"\$\\times\$", "×", text)
    text = re.sub(r"\$\\approx\$", "≈", text)
    text = re.sub(r"\$([^$]*)\$", r"\1", text)   # inline math restante

    # 5. Comandos layout SIN argumento - ANTES de los de formato
    #    (?![a-zA-Z]) evita que \small se fusione con el texto siguiente.
    text = re.sub(
        r"\\(?:small|normalsize|large|Large|huge|noindent|centering|"
        r"raggedright|raggedleft|newline|linebreak|par)(?![a-zA-Z])",
        " ", text
    )
    text = re.sub(r"\\(?:vspace|hspace)\*?\{[^}]*\}", "", text)

    # 6. Formato de texto: multiples pasadas para anidamiento
    #    El patron maneja hasta un nivel de llaves anidadas en el contenido.
    for _ in range(7):
        prev = text
        text = re.sub(
            r"\\(?:textbf|textit|emph|texttt|text|mbox|footnote)"
            r"\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}",
            r"\1", text
        )
        if text == prev:
            break

    # 7. Limpieza final: comandos LaTeX residuales
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{[^}]*\}", "", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?![a-zA-Z])", "", text)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r" {2,}", " ", text)
    text = text.replace("\n", " ")
    return text.strip()


# ── Marcadores internos para formatting inline ──────────────────────────────
BOLD_OPEN  = "\x02B"
BOLD_CLOSE = "\x02b"
ITAL_OPEN  = "\x02I"
ITAL_CLOSE = "\x02i"
CODE_OPEN  = "\x02C"
CODE_CLOSE = "\x02c"


def mark_inline(text: str) -> str:
    """Añade marcadores internos para negrita/cursiva/código."""
    text = re.sub(r"\\textbf\{([^}]+)\}",
                  lambda m: BOLD_OPEN + m.group(1) + BOLD_CLOSE, text)
    text = re.sub(r"\\textit\{([^}]+)\}",
                  lambda m: ITAL_OPEN + m.group(1) + ITAL_CLOSE, text)
    text = re.sub(r"\\emph\{([^}]+)\}",
                  lambda m: ITAL_OPEN + m.group(1) + ITAL_CLOSE, text)
    text = re.sub(r"\\texttt\{([^}]+)\}",
                  lambda m: CODE_OPEN + m.group(1) + CODE_CLOSE, text)
    return text


class Block:
    """Unidad mínima de contenido parseada de LaTeX."""
    def __init__(self, kind: str, **kw):
        self.kind = kind  # heading|para|list|table|figure|caption|source|blank|verbatim
        self.__dict__.update(kw)


def parse_tex_file(path: Path) -> list[Block]:
    """Parsea un archivo .tex y devuelve una lista de bloques."""
    text = path.read_text(encoding="utf-8")
    # Remove LaTeX comments: only % NOT preceded by backslash (i.e., \% is escaped percent).
    # Using negative lookbehind (?<!\\) so that \% survives.
    text = re.sub(r"(?<!\\)%[^\n]*", "", text)
    # Normalize whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks: list[Block] = []
    _parse_body(text, blocks)
    return blocks


def _parse_body(text: str, blocks: list[Block]):
    pos = 0
    n = len(text)

    def peek(s): return text[pos:pos+len(s)] == s

    while pos < n:
        # Skip pure whitespace lines
        m = re.match(r"\n{2,}", text[pos:])
        if m:
            pos += m.end()
            continue

        # Chapter
        m = re.match(r"\\chapter\*?\{([^}]+)\}", text[pos:])
        if m:
            starred = "\\chapter*" in text[pos:pos+10]
            blocks.append(Block("heading", level=0 if not starred else -1,
                                text=clean_inline(m.group(1))))
            pos += m.end(); continue

        # Section
        m = re.match(r"\\section\*?\{([^}]+)\}", text[pos:])
        if m:
            blocks.append(Block("heading", level=1, text=clean_inline(m.group(1))))
            pos += m.end(); continue

        # Subsection
        m = re.match(r"\\subsection\*?\{([^}]+)\}", text[pos:])
        if m:
            blocks.append(Block("heading", level=2, text=clean_inline(m.group(1))))
            pos += m.end(); continue

        # Table environment
        m = re.match(r"\\begin\{table\}.*?\\end\{table\}", text[pos:], re.DOTALL)
        if m:
            _parse_table_env(m.group(0), blocks)
            pos += m.end(); continue

        # Longtable environment
        m = re.match(r"\\begin\{longtable\}.*?\\end\{longtable\}", text[pos:], re.DOTALL)
        if m:
            _parse_table_env(m.group(0), blocks, longtable=True)
            pos += m.end(); continue

        # Figure environment
        m = re.match(r"\\begin\{figure\}.*?\\end\{figure\}", text[pos:], re.DOTALL)
        if m:
            _parse_figure_env(m.group(0), blocks)
            pos += m.end(); continue

        # itemize / enumerate
        m = re.match(r"\\begin\{(itemize|enumerate)\}(.*?)\\end\{\1\}", text[pos:], re.DOTALL)
        if m:
            numbered = m.group(1) == "enumerate"
            items = re.findall(r"\\item\s*(.*?)(?=\\item|\\end\{(?:itemize|enumerate)\})",
                               m.group(2), re.DOTALL)
            for i, item in enumerate(items):
                item_text = clean_inline(mark_inline(item.strip()))
                blocks.append(Block("list", text=item_text, numbered=numbered, index=i+1))
            pos += m.end(); continue

        # quote / begin{quote}
        m = re.match(r"\\begin\{quote\}(.*?)\\end\{quote\}", text[pos:], re.DOTALL)
        if m:
            blocks.append(Block("quote", text=clean_inline(mark_inline(m.group(1).strip()))))
            pos += m.end(); continue

        # Skip other environments
        m = re.match(r"\\begin\{[^}]+\}.*?\\end\{[^}]+\}", text[pos:], re.DOTALL)
        if m:
            pos += m.end(); continue

        # Skip LaTeX preamble commands at start of line
        m = re.match(r"\\(?:label|ref|vspace|hspace|clearpage|newpage|pagebreak|phantom"
                     r"|addcontentsline|phantomsection|appendix|setcounter|pagenumbering"
                     r"|normalsize|small|large|Large|noindent|centering|chapter\*?|section\*?"
                     r"|subsection\*?)[^\n]*\n?", text[pos:])
        if m:
            pos += m.end(); continue

        # Regular paragraph text: collect until double newline or known command
        m = re.match(r"([^\n]+(?:\n(?!\n|\\(?:chapter|section|subsection|begin|end|item|label)).*)*)",
                     text[pos:])
        if m:
            raw = m.group(0).strip()
            if raw:
                marked = mark_inline(raw)
                clean = clean_inline(marked)
                if clean and len(clean) > 3:
                    blocks.append(Block("para", text=clean, raw=raw))
            pos += m.end(); continue

        pos += 1  # advance if nothing matched


def _skip_nested_braces(s: str, start: int = 0) -> int:
    """Devuelve el índice del } de cierre que corresponde al { en start."""
    depth = 0
    for i in range(start, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return i
    return len(s) - 1


def _extract_tabular_body(env: str) -> str:
    """Extrae el cuerpo del entorno tabular saltando correctamente la
    especificación de columnas (que puede contener llaves anidadas)."""
    m = re.search(r"\\begin\{(?:tabular|longtable)\}", env)
    if not m:
        return ""
    rest = env[m.end():]
    # La especificación de columnas empieza con {
    if rest and rest[0] == "{":
        end_idx = _skip_nested_braces(rest, 0)
        return rest[end_idx + 1 :]
    return rest


def _extract_caption_text(env: str) -> str:
    """Extrae el texto de \\caption{...} manejando llaves anidadas."""
    m = re.search(r"\\caption", env)
    if not m:
        return ""
    rest = env[m.end():]
    # Puede haber [ ] opcional primero
    rest = re.sub(r"^\[[^\]]*\]", "", rest.lstrip())
    if rest and rest[0] == "{":
        end_idx = _skip_nested_braces(rest, 0)
        return clean_inline(rest[1:end_idx])
    return ""


def _parse_table_env(env: str, blocks: list[Block], longtable: bool = False):
    """Extrae caption, datos y fuente de un entorno table/longtable."""
    caption = _extract_caption_text(env)

    # Source note - usar funcion robusta con soporte de llaves anidadas
    source = _extract_source_note(env)

    # Extraer cuerpo del tabular con manejo correcto de llaves anidadas
    body = _extract_tabular_body(env)
    if not body.strip():
        if caption:
            blocks.append(Block("caption", text=caption))
        return

    # Limpiar encabezados repetidos en longtable
    body = re.sub(r"\\endfirsthead.*?\\endhead", "", body, flags=re.DOTALL)
    body = re.sub(r"\\endfoot.*?", "", body)
    body = re.sub(r"\\endlastfoot.*?", "", body)
    # Quitar caption y source que puedan estar dentro del body
    body = re.sub(r"\\caption\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}", "", body)
    body = re.sub(r"\\small\\textit\{[^}]*\}", "", body)
    body = re.sub(r"\\label\{[^}]*\}", "", body)
    body = re.sub(r"\\vspace\{[^}]*\}", "", body)
    body = re.sub(r"\\normalsize", "", body)

    # Dividir filas por \\ o \tabularnewline
    rows_raw = re.split(r"\\\\|\\tabularnewline", body)
    rows = []
    for row in rows_raw:
        row = row.strip()
        if not row:
            continue
        # Eliminar comandos de línea
        row = re.sub(r"\\(?:toprule|midrule|bottomrule|hline)\b\s*", "", row)
        row = re.sub(r"\\cline\{[^}]+\}\s*", "", row)
        row = re.sub(r"\\endfirsthead|\\endhead|\\endfoot|\\endlastfoot", "", row)
        row = row.strip()
        if not row or set(row) <= {" ", "\n", "\t"}:
            continue
        # Dividir celdas por &
        cells = [clean_inline(c.strip()) for c in row.split("&")]
        # Filtrar filas que solo contienen residuos del col spec
        if any(c and not c.startswith("p{") and not c.startswith(">") for c in cells):
            rows.append(cells)

    if rows:
        blocks.append(Block("table", caption=caption, source=source, rows=rows))


def _extract_source_note(env: str) -> str:
    """Extrae la nota de fuente de un entorno tabla o figura.

    Maneja llaves anidadas correctamente (por ejemplo, citas dentro de la fuente).
    Busca \\small\\textit{ o \\textit{Fuente usando .find() sin regex.
    """
    idx = -1

    # Search for \small\textit{ (the most common pattern)
    candidate = env.find(r"\small\textit{")
    if candidate != -1:
        idx = candidate + len(r"\small\textit{") - 1  # point to the {

    # Search for \textit{Fuente
    if idx == -1:
        candidate = env.find(r"\textit{Fuente")
        if candidate != -1:
            idx = candidate + len(r"\textit{") - 1  # point to the {

    if idx == -1:
        # Last resort: find any "Fuente" text
        fuente_idx = env.find("Fuente")
        if fuente_idx == -1:
            return ""
        # Find the enclosing opening brace
        brace_start = env.rfind("{", 0, fuente_idx)
        if brace_start == -1:
            end_brace = env.find("}", fuente_idx)
            if end_brace == -1:
                return ""
            return clean_inline(env[fuente_idx:end_brace + 1])
        end = _skip_nested_braces(env[brace_start:], 0)
        inner = env[brace_start + 1:brace_start + end]
        return clean_inline(inner)

    # idx points to the { that starts the textit content
    brace_pos = env.find("{", idx)
    if brace_pos == -1:
        return ""
    end = _skip_nested_braces(env[brace_pos:], 0)
    inner = env[brace_pos + 1:brace_pos + end]
    return clean_inline(inner)


def _parse_figure_env(env: str, blocks: list[Block]):
    """Extrae caption, fuente e imagen de un entorno figure."""
    # Usar _extract_caption_text que maneja llaves anidadas correctamente
    caption = _extract_caption_text(env)
    
    # Nota de fuente con manejo de llaves anidadas
    source = _extract_source_note(env)

    img_path = ""
    # Extract path from \includegraphics[...]{path}
    img_m = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", env)
    if img_m:
        img_path = img_m.group(1)

    if img_path:
        # Resolve path relative to project root
        img_path = img_path.replace("../results/figures/", "results/figures/")
        img_path = img_path.replace("../Memoria/figuras/", "Memoria/figuras/")
        img_path = img_path.replace("figuras/", "Memoria/figuras/")
        full_path = ROOT / img_path
        if not full_path.exists():
            # try alternatives
            for alt in [ROOT / "results" / "figures" / Path(img_path).name,
                        ROOT / "Memoria" / "figuras" / Path(img_path).name]:
                if alt.exists():
                    full_path = alt
                    break
        blocks.append(Block("figure", caption=caption, source=source,
                            img_path=str(full_path)))

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

class DocBuilder:
    """Construye el documento Word aplicando estilos UNIR correctos."""

    def __init__(self, template_path: Path):
        # Abre la plantilla DIRECTAMENTE; luego borramos su contenido.
        self.doc = Document(str(template_path))
        self._clear_body()
        self._fix_headers_footers()

    def _clear_body(self):
        """Elimina TODO el contenido del template, preserva solo el sectPr final."""
        body = self.doc.element.body
        children = list(body)
        # Guardar el ultimo sectPr (contiene margenes, encabezado y pie de pagina)
        last_sectPr = None
        for child in reversed(children):
            if child.tag.split("}")[-1] == "sectPr":
                last_sectPr = child
                break
        # Eliminar TODOS los hijos (incluyendo tablas residuales de la plantilla)
        for child in children:
            body.remove(child)
        # Restaurar sectPr
        if last_sectPr is not None:
            body.append(last_sectPr)

    def _fix_headers_footers(self):
        """Actualiza el encabezado con los nombres correctos de los autores."""
        authors = "Rico Ballesteros, J. C.  ·  Blanco Mendoza, J. J.  ·  Machado Loaiza, J. M."
        title_short = "Evaluación de modelos de pronóstico para FP&A"
        for sec in self.doc.sections:
            # Header
            if sec.header:
                for para in sec.header.paragraphs:
                    if para.text.strip():
                        # Left part = authors, Right part = title
                        para.clear()
                        run = para.add_run(f"{authors}     {title_short}")
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
                        break

    # ── Helpers ──────────────────────────────────────────────────────────────

    def _add_para(self, text: str, style: str = "Normal",
                  align: WD_ALIGN_PARAGRAPH = None) -> docx.text.paragraph.Paragraph:
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.paragraph_format.alignment = align
        _add_runs(p, text)
        return p

    def _add_table(self, rows: list[list[str]], style: str = "Table Grid",
                   has_header: bool = True) -> docx.table.Table:
        if not rows:
            return None
        n_cols = max(len(r) for r in rows)
        n_cols = max(n_cols, 1)
        tbl = self.doc.add_table(rows=len(rows), cols=n_cols)
        tbl.style = "Table Grid"
        tbl.autofit = True
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j >= n_cols:
                    break
                cell = tbl.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.style = self.doc.styles["Normal"]
                _add_runs(p, cell_text)
                if has_header and i == 0:
                    for run in p.runs:
                        run.bold = True
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
        return tbl

    # ── Cover page ────────────────────────────────────────────────────────────

    def add_cover(self):
        # Universidad
        p = self.doc.add_paragraph(style="No Spacing")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("Universidad Internacional de La Rioja")
        r.font.name = "Calibri"; r.font.size = Pt(16)

        p = self.doc.add_paragraph(style="No Spacing")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("Escuela Superior de Ingeniería y Tecnología")
        r.font.name = "Calibri"; r.font.size = Pt(14)

        p = self.doc.add_paragraph(style="No Spacing")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("Máster Universitario en Inteligencia Artificial")
        r.font.name = "Calibri"; r.font.size = Pt(13)

        # Espacio
        for _ in range(4):
            self.doc.add_paragraph(style="No Spacing")

        # Título
        p = self.doc.add_paragraph(style="No Spacing")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(
            "Evaluación de modelos de pronóstico para la planificación "
            "financiera (FP&A) mediante backtesting temporal y métricas "
            "de impacto económico"
        )
        r.font.name = "Calibri Light"; r.font.size = Pt(22)
        r.font.bold = True
        r.font.color.rgb = UNIR_BLUE

        # Espacio
        for _ in range(5):
            self.doc.add_paragraph(style="No Spacing")

        # Datos del trabajo
        datos = [
            ("Autores:", "Juan Camilo Rico Ballesteros"),
            ("",          "Juan José Blanco Mendoza"),
            ("",          "José Manuel Machado Loaiza"),
            ("Tipo de trabajo:", "Comparativa de soluciones"),
            ("Directora:", "Marta María Arguedas Lafuente"),
            ("Curso académico:", "2024-2025"),
        ]
        for etiqueta, valor in datos:
            p = self.doc.add_paragraph(style="No Spacing")
            if etiqueta:
                r = p.add_run(f"{etiqueta}  ")
                r.font.name = "Calibri"; r.font.size = Pt(12); r.bold = True
            r = p.add_run(valor)
            r.font.name = "Calibri"; r.font.size = Pt(12)

        self.doc.add_page_break()

    # ── Section title (Título Índices) ────────────────────────────────────────

    def add_titulo_indice(self, text: str):
        self._add_para(text, style="Título Índices")

    # ── Headings ──────────────────────────────────────────────────────────────

    def add_heading(self, text: str, level: int):
        """level 0=chapter(H1), 1=section(H2), 2=subsection(H3), -1=unnumbered.

        Las subsecciones (Heading 3) se renderizan como parrafo Normal en negrita
        sin numeracion, siguiendo la recomendacion del revisor de evitar
        subsecciones numeradas en la metodologia.
        """
        if level == 2:
            # Subsecciones: negrita sin numeracion automatica
            p = self.doc.add_paragraph(style="Normal")
            _add_runs(p, text)
            for run in p.runs:
                run.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            return p
        if level == -1:
            style = "Título 1 sin numerar"
        elif level == 0:
            style = "Heading 1"
        else:
            style = "Heading 2"
        p = self.doc.add_paragraph(style=style)
        _add_runs(p, text)
        return p

    # ── Content blocks ────────────────────────────────────────────────────────

    def add_block(self, block: Block):
        if block.kind == "heading":
            if block.level == 0:  # Salto de pagina antes de cada capitulo
                self.doc.add_page_break()
            self.add_heading(block.text, block.level)

        elif block.kind == "para":
            self._add_para(block.text, style="Normal")

        elif block.kind == "quote":
            self._add_para(block.text, style="Normal",
                           align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        elif block.kind == "list":
            p = self.doc.add_paragraph(style="List Paragraph")
            if block.numbered:
                p.paragraph_format.left_indent = Cm(1.27)
                run = p.add_run(f"{block.index}. ")
                run.bold = False
            else:
                # Bullet via em-dash
                run = p.add_run("• ")
            _add_runs(p, block.text)

        elif block.kind == "table":
            self._add_table_block(block)

        elif block.kind == "figure":
            self._add_figure_block(block)

        elif block.kind == "caption":
            self._add_para(block.text, style="Caption")

    def _add_table_block(self, block: Block):
        if block.caption:
            self._add_para(block.caption, style="Caption")
        self._add_table(block.rows)
        if block.source:
            src = block.source
            src = re.sub(r"\\small\\textit\{([^}]+)\}", r"\1", src)
            src = clean_inline(src)
            self._add_para(src, style="Pie de foto-tabla")

    def _add_figure_block(self, block: Block):
        # Caption ABOVE (UNIR standard)
        if block.caption:
            self._add_para(block.caption, style="Caption")

        # Image
        img_path = block.img_path
        if img_path and Path(img_path).exists():
            p = self.doc.add_paragraph(style="Figuras")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(img_path, width=Cm(14.5))
            except Exception as e:
                p.add_run(f"[Figura: {Path(img_path).name}]")
        else:
            self._add_para(f"[Figura no encontrada: {img_path}]", "Normal")

        # Source BELOW
        if block.source:
            src = re.sub(r"\\small\\textit\{([^}]+)\}", r"\1", block.source)
            src = clean_inline(src)
            self._add_para(src, style="Pie de foto-tabla")

    # ── Bibliography ──────────────────────────────────────────────────────────

    def add_references(self):
        self.add_heading("Referencias bibliográficas", -1)
        for _, ref_text in REFERENCIAS_APA:
            p = self.doc.add_paragraph(style="Referencias bibliográficas")
            _add_runs(p, ref_text)

    # ── Save ──────────────────────────────────────────────────────────────────

    def _insert_word_field(self, field_code: str):
        """Inserta un campo Word (TOC, LoF, LoT) actualizable con F9 en Word."""
        from docx.oxml import OxmlElement as El
        from docx.oxml.ns import qn
        p = self.doc.add_paragraph(style="Normal")
        run = p.add_run()
        fc_begin = El("w:fldChar")
        fc_begin.set(qn("w:fldCharType"), "begin")
        instr = El("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_code
        fc_sep = El("w:fldChar")
        fc_sep.set(qn("w:fldCharType"), "separate")
        run2 = p.add_run("[Actualizar con clic derecho → Actualizar campo]")
        run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run2.font.italic = True
        fc_end = El("w:fldChar")
        fc_end.set(qn("w:fldCharType"), "end")
        run3 = p.add_run()
        run._r.append(fc_begin)
        run._r.append(instr)
        run._r.append(fc_sep)
        run3._r.append(fc_end)
        return p

    def add_indices_pages(self):
        """Inserta las paginas de Indice de contenidos, figuras y tablas."""
        self.add_titulo_indice("Índice de contenidos")
        self._insert_word_field('TOC \\o "1-3" \\h \\z \\u')
        self.doc.add_page_break()
        self.add_titulo_indice("Índice de figuras")
        self._insert_word_field('TOC \\h \\z \\c "Figuras"')
        self.doc.add_page_break()
        self.add_titulo_indice("Índice de tablas")
        self._insert_word_field('TOC \\h \\z \\c "Tabla"')
        self.doc.add_page_break()

    def save(self, path: Path):
        self.doc.save(str(path))
        print(f"Guardado: {path}")


def _add_runs(para: docx.text.paragraph.Paragraph, text: str):
    """Añade runs con negrita/cursiva/código según los marcadores internos."""
    # Tokenize: split on markers
    tokens = re.split(
        f"({re.escape(BOLD_OPEN)}|{re.escape(BOLD_CLOSE)}"
        f"|{re.escape(ITAL_OPEN)}|{re.escape(ITAL_CLOSE)}"
        f"|{re.escape(CODE_OPEN)}|{re.escape(CODE_CLOSE)})",
        text
    )
    bold = italic = code = False
    for token in tokens:
        if token == BOLD_OPEN: bold = True
        elif token == BOLD_CLOSE: bold = False
        elif token == ITAL_OPEN: italic = True
        elif token == ITAL_CLOSE: italic = False
        elif token == CODE_OPEN: code = True
        elif token == CODE_CLOSE: code = False
        elif token:
            run = para.add_run(token)
            if bold: run.bold = True
            if italic: run.italic = True
            if code:
                run.font.name = "Courier New"
                run.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER FILES TO PROCESS
# ═══════════════════════════════════════════════════════════════════════════════

CAPS_DIR = MEMORIA / "capitulos"

CHAPTER_FILES = [
    ("resumen",       CAPS_DIR / "resumen.tex",           "front"),
    ("organizacion",  CAPS_DIR / "organizacion.tex",      "front"),
    ("cap1",          CAPS_DIR / "cap1_introduccion.tex", "body"),
    ("cap2",          CAPS_DIR / "cap2_contexto.tex",     "body"),
    ("cap3",          CAPS_DIR / "cap3_objetivos.tex",    "body"),
    ("cap4",          CAPS_DIR / "cap4_planteamiento.tex","body"),
    ("cap5",          CAPS_DIR / "cap5_resultados.tex",   "body"),
    ("cap6",          CAPS_DIR / "cap6_discusion.tex",    "body"),
    ("cap7",          CAPS_DIR / "cap7_conclusiones.tex", "body"),
    ("anexo",         CAPS_DIR / "anexo_a.tex",           "back"),
]


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print(f"Plantilla: {TEMPLATE}")
    print(f"Salida:    {OUTPUT}")

    builder = DocBuilder(TEMPLATE)

    # ── Portada ───────────────────────────────────────────────────────────────
    builder.add_cover()

    # ── Índices (ToC, LoF, LoT) ────────────────────────────────────────────────
    builder.add_indices_pages()

    # ── Resumen / Abstract / Organización ─────────────────────────────────────
    for name, path, section_type in CHAPTER_FILES:
        if not path.exists():
            print(f"  SKIP (no existe): {path.name}")
            continue
        print(f"  Procesando: {path.name}")
        blocks = parse_tex_file(path)
        for block in blocks:
            # Ajustar heading level para secciones sin numerar del front matter
            if section_type == "front" and block.kind == "heading" and block.level == 0:
                block.level = -1  # → Título Índices
            builder.add_block(block)

        if section_type == "front" and name == "organizacion":
            builder.doc.add_page_break()

    # ── Referencias bibliográficas ─────────────────────────────────────────────
    builder.add_references()

    # ── Guardar ───────────────────────────────────────────────────────────────
    builder.save(OUTPUT)


if __name__ == "__main__":
    main()
