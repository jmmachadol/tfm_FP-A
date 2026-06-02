"""Exporta la memoria LaTeX a formato Word (.docx).

Usa pandoc (si está instalado) o, como fallback, genera un .docx con la
librería python-docx a partir del contenido de los capítulos .tex.

Uso:
    python export_to_word.py                    # exporta a memoria/TFM_Entrega2.docx
    python export_to_word.py --output mi.docx   # destino personalizado
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).parent
MEMORIA_DIR = ROOT / "memoria"
DEFAULT_OUT = MEMORIA_DIR / "TFM_Entrega2.docx"


def export_with_pandoc(output: Path) -> bool:
    """Intenta exportar usando pandoc. Devuelve True si tiene éxito."""
    pandoc = shutil.which("pandoc")
    if not pandoc:
        # Intentar instalación con winget
        print("pandoc no encontrado en PATH. Intentando instalar con winget...")
        try:
            subprocess.run(
                ["winget", "install", "JohnMacFarlane.Pandoc",
                 "--silent", "--accept-package-agreements", "--accept-source-agreements"],
                check=True, capture_output=True
            )
            pandoc = shutil.which("pandoc")
        except Exception:
            pass

    if not pandoc:
        return False

    print(f"Usando pandoc: {pandoc}")
    cmd = [
        pandoc,
        str(MEMORIA_DIR / "main.tex"),
        "--from=latex+raw_tex",
        "--to=docx",
        f"--output={output}",
        "--reference-doc=" + str(MEMORIA_DIR / "reference.docx") if (MEMORIA_DIR / "reference.docx").exists() else "",
        "--citeproc",
        f"--bibliography={MEMORIA_DIR / 'referencias.bib'}",
        "--csl=apa.csl",
        "--wrap=none",
    ]
    cmd = [c for c in cmd if c]  # eliminar vacíos
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(MEMORIA_DIR))
        if result.returncode == 0:
            print(f"✓ Exportado a: {output}")
            return True
        else:
            print(f"pandoc error: {result.stderr[:500]}")
            return False
    except Exception as e:
        print(f"pandoc falló: {e}")
        return False


def collect_tex_content() -> str:
    """Lee y concatena el contenido de todos los capítulos .tex."""
    chapter_files = [
        "capitulos/resumen.tex",
        "capitulos/organizacion.tex",
        "capitulos/cap1_introduccion.tex",
        "capitulos/cap2_contexto.tex",
        "capitulos/cap3_objetivos.tex",
        "capitulos/cap4_planteamiento.tex",
        "capitulos/cap5_resultados.tex",
        "capitulos/cap6_discusion.tex",
        "capitulos/cap7_conclusiones.tex",
        "capitulos/anexo_a.tex",
    ]
    content = ""
    for fname in chapter_files:
        path = MEMORIA_DIR / fname
        if path.exists():
            content += f"\n\n% ===== {fname} =====\n"
            content += path.read_text(encoding="utf-8")
    return content


def clean_latex(text: str) -> str:
    """Limpieza básica de marcado LaTeX para obtener texto plano."""
    # Eliminar comentarios
    text = re.sub(r"%.*", "", text)
    # Eliminar comandos de sección preservando el título
    for cmd in ["chapter", "section", "subsection", "subsubsection"]:
        text = re.sub(rf"\\{cmd}\*?\{{([^}}]+)\}}", r"\1\n" + "=" * 40 + "\n", text)
    # Comandos simples a texto
    text = re.sub(r"\\textbf\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\label\{[^}]+\}", "", text)
    text = re.sub(r"\\ref\{[^}]+\}", "[ref]", text)
    text = re.sub(r"\\parencite\{([^}]+)\}", r"(\1)", text)
    text = re.sub(r"\\textcite\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\cite\{([^}]+)\}", r"(\1)", text)
    text = re.sub(r"\\url\{([^}]+)\}", r"\1", text)
    # Entornos
    text = re.sub(r"\\begin\{[^}]+\}", "", text)
    text = re.sub(r"\\end\{[^}]+\}", "", text)
    # Caracteres especiales LaTeX
    text = text.replace("\\\\", "\n").replace("\\&", "&").replace("\\%", "%")
    text = text.replace("\\,", " ").replace("~", " ").replace("---", "—").replace("--", "–")
    # Eliminar comandos restantes
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})*", " ", text)
    text = re.sub(r"\{|\}", "", text)
    # Limpiar espacios múltiples
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def export_with_python_docx(output: Path) -> bool:
    """Genera un .docx básico con python-docx como fallback."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        try:
            subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"], check=True)
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except Exception as e:
            print(f"python-docx no disponible: {e}")
            return False

    print("Generando .docx con python-docx...")
    doc = Document()

    # Configurar página A4 con márgenes UNIR
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    BLUE = RGBColor(0, 70, 127)

    # Portada
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Evaluación de modelos de pronóstico para la planificación financiera (FP&A) "
        "mediante backtesting temporal y métricas de impacto económico"
    )
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE
    run.bold = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Universidad Internacional de La Rioja (UNIR)\n"
                 "Máster Universitario en Inteligencia Artificial\n\n"
                 "Autores: Juan Camilo Rico Ballesteros · Juan José Blanco Mendoza · "
                 "José Manuel Machado Loaiza\n"
                 "Directora: Marta María Arguedas Lafuente")

    doc.add_page_break()

    # Leer y limpiar capítulos
    raw = collect_tex_content()
    clean = clean_latex(raw)

    # Procesar capítulos identificados por "=========="
    current_chapter = None
    para_buffer = []

    def flush_buffer(d, buf, level):
        if buf:
            text = " ".join(buf).strip()
            if text:
                p = d.add_paragraph(text)
                p.style = "Body Text"
            buf.clear()

    lines = clean.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "=" * 10 in line:
            flush_buffer(doc, para_buffer, 1)
            # El texto antes de la línea de = es el título
            if i > 0:
                title_line = lines[i - 1].strip() if i > 0 else ""
                if title_line and "capitulos/" not in title_line:
                    h = doc.add_heading(title_line, level=1)
                    for run in h.runs:
                        run.font.color.rgb = BLUE
        elif line and not line.startswith("%"):
            para_buffer.append(line)
            if line.endswith(".") or line.endswith(":"):
                flush_buffer(doc, para_buffer, 1)
        i += 1

    flush_buffer(doc, para_buffer, 1)

    doc.save(str(output))
    print(f"✓ Exportado a: {output}")
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description="Exportar memoria LaTeX a Word")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT,
                        help="Ruta del archivo .docx de salida")
    args = parser.parse_args()

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    print(f"Exportando memoria LaTeX → Word: {output}")

    # Intentar pandoc primero (mejor calidad)
    if export_with_pandoc(output):
        return

    # Fallback: python-docx
    print("Usando python-docx como fallback...")
    if export_with_python_docx(output):
        return

    print("ERROR: No se pudo exportar. Instale pandoc o python-docx.", file=sys.stderr)
    sys.exit(1)


if __name__ == "__main__":
    main()
